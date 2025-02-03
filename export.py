# Program je šířen zdarma. Můžete jej upravovat i šířit dále.
# Tento program je poskytován „tak, jak je“, bez jakýchkoliv záruk,
# výslovných či předpokládaných, včetně (mimo jiné) předpokládaných
# záruk prodejnosti nebo vhodnosti pro určitý účel. Autor nenese
# odpovědnost za jakékoli přímé, nepřímé nebo následné škody vyplývající
# z použití tohoto programu. Používáte jej výhradně na vlastní nebezpečí.

# Martin Kolba - 2024





import os
import logging
import io
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, scrolledtext
import threading
from docx import Document
from PIL import Image
import fitz  # PyMuPDF
import traceback

# Nastavení logování – ukládá se do souboru extraction_log.txt
logging.basicConfig(
    filename="extraction_log.txt",
    level=logging.ERROR,
    format="%(asctime)s - %(levelname)s - %(message)s"
)


def save_images(document, image_folder, image_format='PNG', image_quality=90):
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)

    # Procházení vztahů dokumentu a uložení obrázků
    for i, rel in enumerate(document.part.rels.values()):
        if "image" in rel.target_ref:
            try:
                img_part = rel.target_part
                image_data = img_part.blob
                image = Image.open(io.BytesIO(image_data))
                image_path = os.path.join(image_folder, f'image_{i+1:03}.{image_format.lower()}')
                image.save(image_path, format=image_format, quality=image_quality)
            except Exception as e:
                logging.error(f"Chyba při ukládání obrázku č.{i+1}: {e}")
                continue


def extract_text_docx(document, text_file, preserve_styles):
    # Vytvoří cílovou složku, pokud je definována
    output_dir = os.path.dirname(text_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    new_doc = Document()
    extracted_text = ""

    if preserve_styles:
        for paragraph in document.paragraphs:
            new_paragraph = new_doc.add_paragraph()
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                extracted_text += run.text
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
            new_paragraph.style = paragraph.style
            new_doc.add_paragraph()  # Oddělovač mezi odstavci
    else:
        for paragraph in document.paragraphs:
            extracted_text += paragraph.text + "\n"
            new_doc.add_paragraph(paragraph.text)

    new_doc.save(text_file)
    return extracted_text


def extract_text_pdf(pdf_file, text_file):
    output_dir = os.path.dirname(text_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    extracted_text = ""
    try:
        with fitz.open(pdf_file) as doc:
            for page in doc:
                extracted_text += page.get_text() + "\n"
    except Exception as e:
        logging.error(f"Chyba při čtení PDF: {e}")
        raise

    # Uložíme text do DOCX – vzhledem k absenci stylů u PDF se jedná o "plain" verzi
    docx_doc = Document()
    docx_doc.add_paragraph(extracted_text)
    docx_doc.save(text_file)
    return extracted_text


def extract_images_pdf(pdf_file, image_folder, image_format='PNG', image_quality=90):
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)

    try:
        doc = fitz.open(pdf_file)
        for page_number in range(len(doc)):
            page = doc.load_page(page_number)
            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes))
                    image_path = os.path.join(
                        image_folder,
                        f'image_{page_number+1}_{img_index+1:03}.{image_format.lower()}'
                    )
                    image.save(image_path, format=image_format, quality=image_quality)
                except Exception as e:
                    logging.error(f"Chyba při ukládání obrázku z PDF na stránce {page_number+1}: {e}")
                    continue
    except Exception as e:
        logging.error(f"Chyba při otevírání PDF souboru: {e}")
        raise


def save_preview(text):
    """Umožní uživateli uložit náhled textu do souboru."""
    save_path = filedialog.asksaveasfilename(
        title="Uložit náhled do souboru",
        defaultextension=".txt",
        filetypes=[("Textové soubory", "*.txt"), ("Všechny soubory", "*.*")]
    )
    if save_path:
        try:
            with open(save_path, "w", encoding="utf-8") as f:
                f.write(text)
            messagebox.showinfo("Uloženo", f"Náhled byl uložen do:\n{save_path}")
        except Exception as e:
            logging.error(f"Chyba při ukládání náhledu: {e}")
            messagebox.showerror("Chyba", f"Došlo k chybě při ukládání náhledu: {e}")


def show_preview(text):
    """Zobrazí náhled extrahovaného textu v novém okně."""
    preview_window = tk.Toplevel(root)
    preview_window.title("Náhled textu")
    preview_text = scrolledtext.ScrolledText(preview_window, wrap=tk.WORD, width=80, height=20)
    preview_text.insert(tk.END, text)
    preview_text.config(state=tk.DISABLED)
    preview_text.pack(padx=10, pady=10)
    tk.Button(preview_window, text="Uložit náhled", command=lambda: save_preview(text)).pack(pady=10)


def process_files():
    # Získání hodnot z formuláře
    file_path = file_path_var.get()
    output_folder = output_folder_var.get()
    image_format = image_format_var.get()
    try:
        image_quality = int(image_quality_var.get())
    except ValueError:
        root.after(0, lambda: messagebox.showerror("Chyba", "Kvalita obrázků musí být číslo."))
        progress_bar.stop()
        return

    if not file_path or not output_folder:
        root.after(0, lambda: messagebox.showerror("Chyba", "Prosím, vyberte soubor i cílovou složku."))
        progress_bar.stop()
        return

    # Příprava cest k výstupním souborům
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    text_file_styled = os.path.join(output_folder, f"{base_name}_styled.docx")
    text_file_plain = os.path.join(output_folder, f"{base_name}_plain.docx")
    image_folder = os.path.join(output_folder, "Images")

    try:
        # Nastavení progress baru
        progress_bar.config(mode="indeterminate")
        progress_bar.start()

        extracted_text = ""

        if file_path.lower().endswith((".docx", ".doc")):
            # Upozornění: python-docx plně podporuje pouze formát DOCX
            if file_path.lower().endswith(".doc"):
                root.after(0, lambda: messagebox.showwarning(
                    "Upozornění",
                    "Formát .doc nemusí být plně podporován. Doporučuje se převést soubor do DOCX."
                ))
            document = Document(file_path)
            # Verze se zachováním stylů
            extracted_text = extract_text_docx(document, text_file_styled, preserve_styles=True)
            # Verze prostého textu
            extract_text_docx(document, text_file_plain, preserve_styles=False)
            save_images(document, image_folder, image_format, image_quality)

        elif file_path.lower().endswith(".pdf"):
            # U PDF souborech nelze zachovat styly – obě verze budou stejné.
            extracted_text = extract_text_pdf(file_path, text_file_plain)
            # Pro názvy výstupních souborů u PDF vytvoříme pouze jednu verzi (plain)
            # Pokud si přejete obě verze, lze soubor jednoduše zkopírovat:
            # shutil.copy(text_file_plain, text_file_styled)
            extract_images_pdf(file_path, image_folder, image_format, image_quality)
        else:
            raise ValueError("Nepodporovaný formát souboru.")

        # Zobrazení náhledu a informací o uložení souborů – provádíme v hlavním vlákně
        root.after(0, lambda: show_preview(extracted_text))
        root.after(0, lambda: messagebox.showinfo(
            "Dokončeno",
            f"Text byl uložen do:\n{text_file_styled if file_path.lower().endswith(('.docx','.doc')) else text_file_plain}"
        ))
        if os.path.exists(image_folder):
            root.after(0, lambda: messagebox.showinfo("Informace", f"Obrázky byly uloženy do složky:\n{image_folder}"))

    except Exception as e:
        logging.error("Chyba při zpracování souboru:\n" + traceback.format_exc())
        root.after(0, lambda: messagebox.showerror("Chyba", f"Došlo k chybě při zpracování souboru:\n{e}"))
    finally:
        progress_bar.stop()


def start_processing_thread():
    # Spustíme zpracování v odděleném vlákně, aby nedošlo k zablokování GUI
    threading.Thread(target=process_files, daemon=True).start()


def select_file():
    file_path = filedialog.askopenfilename(
        title="Vyberte soubor",
        filetypes=[
            ("Podporované soubory", "*.docx *.doc *.pdf"),
            ("Word soubory", "*.docx *.doc"),
            ("PDF soubory", "*.pdf")
        ]
    )
    if file_path:
        file_path_var.set(file_path)


def select_output_folder():
    folder_path = filedialog.askdirectory(title="Vyberte cílovou složku")
    if folder_path:
        output_folder_var.set(folder_path)


# Hlavní okno aplikace
root = tk.Tk()
root.title("Extrahování textů a obrázků z dokumentů")
root.geometry("800x500")
root.resizable(False, False)

# Proměnné pro uložení cest a nastavení
file_path_var = tk.StringVar()
output_folder_var = tk.StringVar()
image_format_var = tk.StringVar(value="PNG")
image_quality_var = tk.StringVar(value="100")

# Hlavní rámec
frame_main = ttk.Frame(root, padding="10")
frame_main.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

# Rámec pro výběr souboru
frame_file_selection = ttk.LabelFrame(frame_main, text="Výběr souboru", padding="10")
frame_file_selection.grid(row=0, column=0, padx=10, pady=10, sticky=(tk.W, tk.E))

ttk.Label(frame_file_selection, text="Vyberte soubor:").grid(row=0, column=0, padx=10, pady=5, sticky=tk.W)
ttk.Entry(frame_file_selection, textvariable=file_path_var, width=50).grid(row=0, column=1, padx=10, pady=5)
ttk.Button(frame_file_selection, text="Procházet", command=select_file).grid(row=0, column=2, padx=10, pady=5)

# Rámec pro výběr cílové složky
frame_output_selection = ttk.LabelFrame(frame_main, text="Výběr cílové složky", padding="10")
frame_output_selection.grid(row=1, column=0, padx=10, pady=10, sticky=(tk.W, tk.E))

ttk.Label(frame_output_selection, text="Cílová složka:").grid(row=0, column=0, padx=10, pady=5, sticky=tk.W)
ttk.Entry(frame_output_selection, textvariable=output_folder_var, width=50).grid(row=0, column=1, padx=10, pady=5)
ttk.Button(frame_output_selection, text="Vybrat složku", command=select_output_folder).grid(row=0, column=2, padx=10, pady=5)

# Rámec pro nastavení obrázků
frame_options = ttk.LabelFrame(frame_main, text="Nastavení obrázků", padding="10")
frame_options.grid(row=2, column=0, padx=10, pady=10, sticky=(tk.W, tk.E))

ttk.Label(frame_options, text="Formát obrázků:").grid(row=0, column=0, padx=10, pady=5, sticky=tk.W)
ttk.Combobox(
    frame_options,
    textvariable=image_format_var,
    values=["PNG", "JPEG", "BMP"],
    state="readonly"
).grid(row=0, column=1, padx=10, pady=5)

ttk.Label(frame_options, text="Kvalita obrázků:").grid(row=1, column=0, padx=10, pady=5, sticky=tk.W)
ttk.Entry(frame_options, textvariable=image_quality_var, width=10).grid(row=1, column=1, padx=10, pady=5)

# Rámec pro akce
frame_actions = ttk.Frame(frame_main, padding="10")
frame_actions.grid(row=3, column=0, padx=10, pady=10, sticky=(tk.W, tk.E))
ttk.Button(frame_actions, text="Spustit zpracování", command=start_processing_thread).grid(row=0, column=0, padx=10, pady=5)

# Progress bar
progress_bar = ttk.Progressbar(root, orient="horizontal", length=400, mode="indeterminate")
progress_bar.grid(row=4, column=0, columnspan=3, padx=10, pady=20)

root.mainloop()
